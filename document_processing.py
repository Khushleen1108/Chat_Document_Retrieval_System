import os
import pickle
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from concurrent.futures import ThreadPoolExecutor
from langchain.docstore.document import Document
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from ai_helpers import get_document_embedder
from config import BASE_DOCS_DIR

def handle_file_upload(email, uploaded_files):
    user_upload_dir = os.path.join(BASE_DOCS_DIR, email)
    os.makedirs(user_upload_dir, exist_ok=True)

    for file_path in uploaded_files:
        file_path = file_path.strip()
        if os.path.exists(file_path):
            save_path = os.path.join(user_upload_dir, os.path.basename(file_path))
            with open(file_path, "rb") as src, open(save_path, "wb") as dest:
                dest.write(src.read())
            print(f"Uploaded: {file_path}")
        else:
            print(f"File not found: {file_path}")

def load_documents_from_directory(email):
    user_docs_dir = os.path.join(BASE_DOCS_DIR, email)
    if not os.path.exists(user_docs_dir):
        return []

    raw_documents = []
    with ThreadPoolExecutor() as executor:
        futures = []
        for file in os.listdir(user_docs_dir):
            file_path = os.path.join(user_docs_dir, file)
            if file.endswith('.txt'):
                futures.append(executor.submit(load_txt_file, file_path))
            elif file.endswith('.pdf'):
                futures.append(executor.submit(load_pdf_file, file_path))
            elif file.endswith('.docx'):
                futures.append(executor.submit(load_docx_file, file_path))

        for future in futures:
            raw_documents.extend(future.result())

    return raw_documents

def load_txt_file(file_path):
    loader = TextLoader(file_path)
    return loader.load()

def load_pdf_file(file_path):
    pages = []
    reader = PdfReader(file_path)
    for page in reader.pages:
        content = page.extract_text()
        if not content:
            content = "[NON-TEXTUAL CONTENT]"
        pages.append(Document(page_content=content))
    return pages

def load_docx_file(file_path):
    doc = DocxDocument(file_path)
    content = "\n".join(
        para.text if para.text.strip() else "[NON-TEXTUAL CONTENT]" 
        for para in doc.paragraphs
    )
    return [Document(page_content=content)]

def initialize_vector_store(raw_documents, email, reset=True):
    user_vector_store_path = f"./uploaded_docs/{email}_vectorstore.pkl"

    if reset or not os.path.exists(user_vector_store_path):
        if raw_documents:
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
            documents = text_splitter.split_documents(raw_documents)
            vectorstore = FAISS.from_documents(documents, get_document_embedder())

            with open(user_vector_store_path, "wb") as f:
                pickle.dump(vectorstore, f)
            return vectorstore
        else:
            return None
    else:
        with open(user_vector_store_path, "rb") as f:
            return pickle.load(f)
